"""
文档类型检查和验证模块
支持前后端文档类型检查，确保文档格式正确
"""
import mimetypes
import magic
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
from dataclasses import dataclass
from enum import Enum

from src.utils.logger import get_logger
from config.settings import settings

logger = get_logger("data.document_validator")


class DocumentType(Enum):
    """支持的文档类型枚举"""
    # Office文档
    DOC = "doc"
    DOCX = "docx"
    PPT = "ppt"
    PPTX = "pptx"
    XLS = "xls"
    XLSX = "xlsx"
    
    # PDF文档
    PDF = "pdf"
    PDF_SCANNED = "pdf_scanned"  # 扫描版PDF
    
    # 网页和标记语言
    HTML = "html"
    MD = "md"
    XML = "xml"
    
    # 电子书
    EPUB = "epub"
    
    # 文本文档
    TXT = "txt"
    
    # 邮件
    EML = "eml"
    MSG = "msg"
    
    # 数据文件
    CSV = "csv"
    JSON = "json"
    
    # 图像文件
    JPG = "jpg"
    JPEG = "jpeg"
    PNG = "png"
    TIFF = "tiff"
    TIF = "tif"


@dataclass
class DocumentInfo:
    """文档信息"""
    file_path: Path
    file_type: DocumentType
    mime_type: str
    file_size: int
    is_valid: bool
    error_message: Optional[str] = None
    is_scanned_pdf: bool = False
    has_images: bool = False
    has_tables: bool = False
    has_formulas: bool = False


class DocumentValidator:
    """文档验证器"""
    
    def __init__(self):
        """初始化文档验证器"""
        # 文件扩展名到文档类型的映射
        self.extension_mapping = {
            '.doc': DocumentType.DOC,
            '.docx': DocumentType.DOCX,
            '.ppt': DocumentType.PPT,
            '.pptx': DocumentType.PPTX,
            '.xls': DocumentType.XLS,
            '.xlsx': DocumentType.XLSX,
            '.pdf': DocumentType.PDF,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.md': DocumentType.MD,
            '.markdown': DocumentType.MD,
            '.xml': DocumentType.XML,
            '.epub': DocumentType.EPUB,
            '.txt': DocumentType.TXT,
            '.eml': DocumentType.EML,
            '.msg': DocumentType.MSG,
            '.csv': DocumentType.CSV,
            '.json': DocumentType.JSON,
            '.jpg': DocumentType.JPG,
            '.jpeg': DocumentType.JPEG,
            '.png': DocumentType.PNG,
            '.tiff': DocumentType.TIFF,
            '.tif': DocumentType.TIF,
        }
        
        # MIME类型到文档类型的映射
        self.mime_mapping = {
            'application/msword': DocumentType.DOC,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.DOCX,
            'application/vnd.ms-powerpoint': DocumentType.PPT,
            'application/vnd.openxmlformats-officedocument.presentationml.presentation': DocumentType.PPTX,
            'application/vnd.ms-excel': DocumentType.XLS,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': DocumentType.XLSX,
            'application/pdf': DocumentType.PDF,
            'text/html': DocumentType.HTML,
            'text/markdown': DocumentType.MD,
            'application/xml': DocumentType.XML,
            'text/xml': DocumentType.XML,
            'application/epub+zip': DocumentType.EPUB,
            'text/plain': DocumentType.TXT,
            'message/rfc822': DocumentType.EML,
            'application/vnd.ms-outlook': DocumentType.MSG,
            'text/csv': DocumentType.CSV,
            'application/json': DocumentType.JSON,
            'image/jpeg': DocumentType.JPEG,
            'image/png': DocumentType.PNG,
            'image/tiff': DocumentType.TIFF,
        }
        
        logger.info(f"文档验证器初始化完成，支持 {len(self.extension_mapping)} 种文件类型")
    
    def validate_file(self, file_path: Union[str, Path]) -> DocumentInfo:
        """
        验证单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            DocumentInfo: 文档信息
        """
        file_path = Path(file_path)
        
        try:
            # 检查文件是否存在
            if not file_path.exists():
                return DocumentInfo(
                    file_path=file_path,
                    file_type=None,
                    mime_type="",
                    file_size=0,
                    is_valid=False,
                    error_message="文件不存在"
                )
            
            # 获取文件信息
            file_size = file_path.stat().st_size
            
            # 检查文件大小
            if file_size > settings.MAX_FILE_SIZE:
                return DocumentInfo(
                    file_path=file_path,
                    file_type=None,
                    mime_type="",
                    file_size=file_size,
                    is_valid=False,
                    error_message=f"文件大小超过限制 ({file_size} > {settings.MAX_FILE_SIZE})"
                )
            
            # 检查文件扩展名
            extension = file_path.suffix.lower()
            if extension not in self.extension_mapping:
                return DocumentInfo(
                    file_path=file_path,
                    file_type=None,
                    mime_type="",
                    file_size=file_size,
                    is_valid=False,
                    error_message=f"不支持的文件类型: {extension}"
                )
            
            # 获取MIME类型
            mime_type = self._get_mime_type(file_path)
            
            # 确定文档类型
            doc_type = self._determine_document_type(file_path, mime_type)
            
            # 检查是否为扫描版PDF
            is_scanned_pdf = False
            if doc_type == DocumentType.PDF:
                is_scanned_pdf = self._is_scanned_pdf(file_path)
                if is_scanned_pdf:
                    doc_type = DocumentType.PDF_SCANNED
            
            # 检查文档特征
            has_images, has_tables, has_formulas = self._analyze_document_features(file_path, doc_type)
            
            return DocumentInfo(
                file_path=file_path,
                file_type=doc_type,
                mime_type=mime_type,
                file_size=file_size,
                is_valid=True,
                is_scanned_pdf=is_scanned_pdf,
                has_images=has_images,
                has_tables=has_tables,
                has_formulas=has_formulas
            )
            
        except Exception as e:
            logger.error(f"文件验证失败: {file_path}, 错误: {e}")
            return DocumentInfo(
                file_path=file_path,
                file_type=None,
                mime_type="",
                file_size=0,
                is_valid=False,
                error_message=str(e)
            )
    
    def validate_files(self, file_paths: List[Union[str, Path]]) -> List[DocumentInfo]:
        """
        批量验证文件
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            List[DocumentInfo]: 文档信息列表
        """
        results = []
        for file_path in file_paths:
            result = self.validate_file(file_path)
            results.append(result)
        
        valid_count = sum(1 for r in results if r.is_valid)
        logger.info(f"批量验证完成: {valid_count}/{len(results)} 个文件有效")
        
        return results
    
    def _get_mime_type(self, file_path: Path) -> str:
        """获取文件MIME类型"""
        try:
            # 使用python-magic获取更准确的MIME类型
            mime_type = magic.from_file(str(file_path), mime=True)
            return mime_type
        except Exception:
            # 回退到mimetypes模块
            mime_type, _ = mimetypes.guess_type(str(file_path))
            return mime_type or "application/octet-stream"
    
    def _determine_document_type(self, file_path: Path, mime_type: str) -> DocumentType:
        """确定文档类型"""
        # 优先使用MIME类型
        if mime_type in self.mime_mapping:
            return self.mime_mapping[mime_type]
        
        # 回退到文件扩展名
        extension = file_path.suffix.lower()
        return self.extension_mapping.get(extension)
    
    def _is_scanned_pdf(self, file_path: Path) -> bool:
        """检查PDF是否为扫描版"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(str(file_path))
            
            # 检查前几页是否包含文本
            text_pages = 0
            total_pages = min(5, len(doc))  # 只检查前5页
            
            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text().strip()
                if len(text) > 50:  # 如果页面有足够的文本
                    text_pages += 1
            
            doc.close()
            
            # 如果大部分页面没有文本，认为是扫描版
            return text_pages < total_pages * 0.3
            
        except Exception as e:
            logger.warning(f"检查PDF类型失败: {e}")
            return False
    
    def _analyze_document_features(self, file_path: Path, doc_type: DocumentType) -> Tuple[bool, bool, bool]:
        """分析文档特征"""
        has_images = False
        has_tables = False
        has_formulas = False
        
        try:
            if doc_type in [DocumentType.DOCX]:
                has_images, has_tables, has_formulas = self._analyze_docx_features(file_path)
            elif doc_type in [DocumentType.PPTX]:
                has_images, has_tables, has_formulas = self._analyze_pptx_features(file_path)
            elif doc_type in [DocumentType.XLSX, DocumentType.XLS]:
                has_images, has_tables, has_formulas = self._analyze_excel_features(file_path)
            elif doc_type == DocumentType.PDF:
                has_images, has_tables, has_formulas = self._analyze_pdf_features(file_path)
            elif doc_type in [DocumentType.JPG, DocumentType.JPEG, DocumentType.PNG, DocumentType.TIFF, DocumentType.TIF]:
                has_images = True
                
        except Exception as e:
            logger.warning(f"分析文档特征失败: {file_path}, 错误: {e}")
        
        return has_images, has_tables, has_formulas
    
    def _analyze_docx_features(self, file_path: Path) -> Tuple[bool, bool, bool]:
        """分析DOCX文档特征"""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            
            has_images = False
            has_tables = len(doc.tables) > 0
            has_formulas = False
            
            # 检查图片
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    has_images = True
                    break
            
            # 检查公式（简单检测）
            for paragraph in doc.paragraphs:
                if any(keyword in paragraph.text.lower() for keyword in ['∑', '∫', '√', '∞', '≤', '≥']):
                    has_formulas = True
                    break
            
            return has_images, has_tables, has_formulas
            
        except Exception:
            return False, False, False
    
    def _analyze_pptx_features(self, file_path: Path) -> Tuple[bool, bool, bool]:
        """分析PPTX文档特征"""
        # 简化实现，后续可以扩展
        return True, False, False  # PPT通常包含图片
    
    def _analyze_excel_features(self, file_path: Path) -> Tuple[bool, bool, bool]:
        """分析Excel文档特征"""
        return False, True, True  # Excel通常包含表格和公式
    
    def _analyze_pdf_features(self, file_path: Path) -> Tuple[bool, bool, bool]:
        """分析PDF文档特征"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(str(file_path))
            
            has_images = False
            has_tables = False
            has_formulas = False
            
            # 检查前几页
            for page_num in range(min(3, len(doc))):
                page = doc[page_num]
                
                # 检查图片
                if page.get_images():
                    has_images = True
                
                # 检查表格（简单检测）
                text = page.get_text()
                if text.count('\t') > 10 or text.count('|') > 10:
                    has_tables = True
                
                # 检查公式
                if any(symbol in text for symbol in ['∑', '∫', '√', '∞', '≤', '≥', 'α', 'β', 'γ']):
                    has_formulas = True
            
            doc.close()
            return has_images, has_tables, has_formulas
            
        except Exception:
            return False, False, False


# 全局验证器实例
document_validator = DocumentValidator()
